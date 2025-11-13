import os
import io
import logging
import hashlib
from typing import List, Dict, Optional, Tuple
from datetime import datetime

import PyPDF2
import docx
import pdfplumber
import tiktoken
from sentence_transformers import SentenceTransformer
import numpy as np

# Import config with fallback to environment variables
try:
    from config import (
        MAX_FILE_SIZE, ALLOWED_EXTENSIONS, CHUNK_SIZE, CHUNK_OVERLAP, EMBEDDING_MODEL
    )
except ImportError:
    # Fallback to environment variables for deployment
    MAX_FILE_SIZE = int(os.getenv('MAX_FILE_SIZE', 10 * 1024 * 1024))  # 10MB default
    ALLOWED_EXTENSIONS = os.getenv('ALLOWED_EXTENSIONS', 'pdf,docx,txt').split(',')
    CHUNK_SIZE = int(os.getenv('CHUNK_SIZE', 500))
    CHUNK_OVERLAP = int(os.getenv('CHUNK_OVERLAP', 50))
    EMBEDDING_MODEL = os.getenv('EMBEDDING_MODEL', 'all-MiniLM-L6-v2')

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Enhanced document processing with embeddings and chunking"""
    
    def __init__(self):
        self.embedding_model = None
        self.tokenizer = None
        self._load_models()
    
    def _load_models(self):
        """Load the embedding model and tokenizer"""
        try:
            logger.info(f"Loading embedding model: {EMBEDDING_MODEL}")
            self.embedding_model = SentenceTransformer(EMBEDDING_MODEL)
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
            logger.info("Models loaded successfully")
        except Exception as e:
            logger.error(f"Error loading models: {str(e)}")
            raise e
    
    def validate_file(self, file) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            # Check file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if file_size > MAX_FILE_SIZE:
                return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({MAX_FILE_SIZE} bytes)"
            
            # Check file extension
            filename = file.filename.lower()
            file_extension = filename.split('.')[-1] if '.' in filename else ''
            
            if file_extension not in ALLOWED_EXTENSIONS:
                return False, f"File type '{file_extension}' not allowed. Supported types: {', '.join(ALLOWED_EXTENSIONS)}"
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using multiple methods for reliability"""
        text = ""
        
        try:
            # Method 1: PyPDF2 for simple PDFs
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            # If PyPDF2 didn't extract much text, try pdfplumber
            if len(text.strip()) < 100:
                logger.info("PyPDF2 extracted minimal text, trying pdfplumber")
                pdf_file = io.BytesIO(file_content)
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise e
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise e
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise e
    
    def extract_text(self, file) -> str:
        """Extract text from file based on its type"""
        try:
            file.seek(0)
            file_content = file.read()
            filename = file.filename.lower()
            
            if filename.endswith('.pdf'):
                return self.extract_text_from_pdf(file_content)
            elif filename.endswith(('.docx', '.doc')):
                return self.extract_text_from_docx(file_content)
            elif filename.endswith('.txt'):
                return self.extract_text_from_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {filename}")
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise e
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken"""
        try:
            tokens = self.tokenizer.encode(text)
            return len(tokens)
        except Exception as e:
            logger.error(f"Error counting tokens: {str(e)}")
            return len(text.split())  # Fallback to word count
    
    def chunk_text(self, text: str) -> List[Dict]:
        """Split text into overlapping chunks for better context preservation"""
        try:
            chunks = []
            sentences = text.split('.')
            current_chunk = ""
            current_tokens = 0
            chunk_id = 0
            
            for i, sentence in enumerate(sentences):
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                sentence += "."  # Add back the period
                sentence_tokens = self.count_tokens(sentence)
                
                # If adding this sentence would exceed chunk size, save current chunk
                if current_tokens + sentence_tokens > CHUNK_SIZE and current_chunk:
                    chunk_data = {
                        'id': chunk_id,
                        'text': current_chunk.strip(),
                        'token_count': current_tokens,
                        'start_sentence': chunk_id * 10,  # Approximate
                        'end_sentence': i
                    }
                    chunks.append(chunk_data)
                    
                    # Start new chunk with overlap
                    overlap_sentences = sentences[max(0, i - CHUNK_OVERLAP // 50):i]
                    current_chunk = " ".join(overlap_sentences).strip()
                    if current_chunk:
                        current_chunk += " "
                    current_tokens = self.count_tokens(current_chunk)
                    chunk_id += 1
                
                current_chunk += sentence + " "
                current_tokens += sentence_tokens
            
            # Add the last chunk if it has content
            if current_chunk.strip():
                chunk_data = {
                    'id': chunk_id,
                    'text': current_chunk.strip(),
                    'token_count': current_tokens,
                    'start_sentence': chunk_id * 10,
                    'end_sentence': len(sentences)
                }
                chunks.append(chunk_data)
            
            logger.info(f"Created {len(chunks)} chunks from text")
            return chunks
            
        except Exception as e:
            logger.error(f"Error chunking text: {str(e)}")
            raise e
    
    def generate_embeddings(self, chunks: List[Dict]) -> List[Dict]:
        """Generate embeddings for text chunks"""
        try:
            if not self.embedding_model:
                raise ValueError("Embedding model not loaded")
            
            texts = [chunk['text'] for chunk in chunks]
            
            logger.info(f"Generating embeddings for {len(texts)} chunks")
            embeddings = self.embedding_model.encode(
                texts,
                batch_size=32,
                show_progress_bar=True,
                convert_to_numpy=True
            )
            
            # Add embeddings to chunks
            for i, chunk in enumerate(chunks):
                chunk['embedding'] = embeddings[i].tolist()
                chunk['embedding_model'] = EMBEDDING_MODEL
            
            logger.info("Embeddings generated successfully")
            return chunks
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {str(e)}")
            raise e
    
    def compute_similarity(self, query_embedding: np.ndarray, chunk_embeddings: List[np.ndarray]) -> List[float]:
        """Compute cosine similarity between query and chunk embeddings"""
        try:
            similarities = []
            
            for chunk_embedding in chunk_embeddings:
                # Compute cosine similarity
                dot_product = np.dot(query_embedding, chunk_embedding)
                norm_a = np.linalg.norm(query_embedding)
                norm_b = np.linalg.norm(chunk_embedding)
                
                if norm_a == 0 or norm_b == 0:
                    similarity = 0.0
                else:
                    similarity = dot_product / (norm_a * norm_b)
                
                similarities.append(float(similarity))
            
            return similarities
            
        except Exception as e:
            logger.error(f"Error computing similarity: {str(e)}")
            return [0.0] * len(chunk_embeddings)
    
    def search_similar_chunks(self, query: str, chunks: List[Dict], top_k: int = 5) -> List[Dict]:
        """Find most similar chunks to a query"""
        try:
            if not chunks:
                return []
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0]
            
            # Get chunk embeddings
            chunk_embeddings = [np.array(chunk['embedding']) for chunk in chunks]
            
            # Compute similarities
            similarities = self.compute_similarity(query_embedding, chunk_embeddings)
            
            # Add similarity scores to chunks and sort
            chunks_with_scores = []
            for i, chunk in enumerate(chunks):
                chunk_copy = chunk.copy()
                chunk_copy['similarity_score'] = similarities[i]
                chunks_with_scores.append(chunk_copy)
            
            # Sort by similarity and return top_k
            chunks_with_scores.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            return chunks_with_scores[:top_k]
            
        except Exception as e:
            logger.error(f"Error searching similar chunks: {str(e)}")
            return []
    
    def generate_document_hash(self, content: str) -> str:
        """Generate a hash for document deduplication"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def process_document(self, file, user_id: str) -> Dict:
        """Complete document processing pipeline"""
        try:
            # Validate file
            is_valid, message = self.validate_file(file)
            if not is_valid:
                raise ValueError(message)
            
            # Extract text
            text = self.extract_text(file)
            if not text.strip():
                raise ValueError("No text could be extracted from the document")
            
            # Generate document hash for deduplication
            doc_hash = self.generate_document_hash(text)
            
            # Chunk text
            chunks = self.chunk_text(text)
            if not chunks:
                raise ValueError("Could not create text chunks from document")
            
            # Generate embeddings
            chunks_with_embeddings = self.generate_embeddings(chunks)
            
            # Prepare document metadata
            document_data = {
                'filename': file.filename,
                'user_id': user_id,
                'file_size': len(file.read()),
                'text_content': text,
                'document_hash': doc_hash,
                'total_chunks': len(chunks_with_embeddings),
                'total_tokens': sum(chunk['token_count'] for chunk in chunks_with_embeddings),
                'chunks': chunks_with_embeddings,
                'processed_at': datetime.utcnow(),
                'embedding_model': EMBEDDING_MODEL,
                'processing_status': 'completed'
            }
            
            logger.info(f"Document processed successfully: {file.filename}")
            return document_data
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise e

# Create global instance
document_processor = DocumentProcessor() 