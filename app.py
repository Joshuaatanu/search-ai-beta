# app.py
from flask import Flask, request, jsonify, render_template, redirect, url_for, session, flash, make_response
from flask_cors import CORS
import requests
# from google import genai
from duckduckgo_search import DDGS
import os
from dotenv import load_dotenv
# Import the Gemini client from Google's generative AI library
import google.generativeai as genai
import arxiv
import re
from flask_pymongo import PyMongo
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from bson.objectid import ObjectId
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
import json
from datetime import datetime, timedelta
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import threading
import io
from werkzeug.utils import secure_filename
import time
from PIL import Image
from utils.avatar import save_profile_picture, generate_random_avatar, get_avatar_url
from utils.activity import track_login_activity, get_login_history

# Visualization libraries
import pandas as pd
import networkx as nx
from scholarly import scholarly
from urllib.parse import urlparse, parse_qs
import plotly
import plotly.graph_objects as go
import plotly.express as px
import random  # Added for random simulation in visualizations

# OAuth Libraries
from authlib.integrations.flask_client import OAuth
import google_auth_oauthlib.flow
import google.oauth2.credentials
import google.oauth2.id_token
import google.auth.transport.requests

# Import models
from models import User, SearchHistory, Favorite, Recommendation

# Import methodology analyzer
from methodology_analyzer import analyze_paper_methodology, analyze_papers_methodologies, compare_methodologies, analyze_methodology, MethodologyAnalyzer

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)

# Configure Flask app
app.secret_key = os.getenv("SECRET_KEY", secrets.token_hex(16))
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
app.config["SESSION_TYPE"] = "filesystem"
app.config["SESSION_PERMANENT"] = True

# Get MongoDB URI from environment or use default
mongodb_uri = os.getenv("MONGODB_URI", "mongodb://localhost:27017/sentino")
print(f"Connecting to MongoDB: {mongodb_uri}")

# Set up MongoDB
app.config["MONGO_URI"] = mongodb_uri
mongo = PyMongo(app)

# Verify the connection is working and initialize collections
with app.app_context():
    try:
        # Check if we can list collections
        collections = mongo.db.list_collection_names()
        print(f"Connected to MongoDB. Existing collections: {collections}")
        
        # Ensure required collections exist
        required_collections = ["users", "search_history", "favorites", "password_resets"]
        for collection in required_collections:
            if collection not in collections:
                print(f"Creating collection: {collection}")
                mongo.db.create_collection(collection)
        
        print("MongoDB collections initialized successfully")
    except Exception as e:
        print(f"MongoDB connection or initialization failed: {e}")
        print("WARNING: Application may not function correctly without MongoDB!")

# Set up Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

# Set up OAuth
oauth = OAuth(app)

# GitHub OAuth
github_client_id = os.getenv("GITHUB_CLIENT_ID")
github_client_secret = os.getenv("GITHUB_CLIENT_SECRET")
if github_client_id and github_client_secret:
    oauth.register(
        name="github",
        client_id=github_client_id,
        client_secret=github_client_secret,
        access_token_url="https://github.com/login/oauth/access_token",
        access_token_params=None,
        authorize_url="https://github.com/login/oauth/authorize",
        authorize_params=None,
        api_base_url="https://api.github.com/",
        client_kwargs={"scope": "user:email"},
    )

# Email configuration
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'True').lower() in ('true', 'yes', '1')
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER')

# Password reset token expiration (in seconds)
PASSWORD_RESET_EXPIRATION = 3600  # 1 hour

# Function to send email
def send_email(to, subject, body_html):
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = app.config['MAIL_DEFAULT_SENDER']
    msg['To'] = to
    
    html_part = MIMEText(body_html, 'html')
    msg.attach(html_part)
    
    try:
        server = smtplib.SMTP(app.config['MAIL_SERVER'], app.config['MAIL_PORT'])
        server.ehlo()
        server.starttls()
        server.login(app.config['MAIL_USERNAME'], app.config['MAIL_PASSWORD'])
        server.sendmail(app.config['MAIL_DEFAULT_SENDER'], to, msg.as_string())
        server.close()
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False

@login_manager.user_loader
def load_user(user_id):
    try:
        return User.get_by_id(mongo.db, user_id)
    except Exception as e:
        print(f"Error loading user: {e}")
        return None

# Verify MongoDB is connected before handling requests
@app.before_request
def check_mongodb():
    if request.endpoint in ['login', 'static', 'github_callback', 'google_callback', 'apple_callback']:
        try:
            # Quick connectivity check
            mongo.db.command('ping')
        except Exception as e:
            print(f"MongoDB not available: {e}")
            if not request.path.startswith('/static/'):
                flash("Database connection issue. Some features may not work properly.", "error")

def query_duckduckgo_text(query):
    """
    Uses DuckDuckGo's text search via the duckduckgo_search library to get search results.
    """
    try:
        ddgs = DDGS()
        results = ddgs.text(keywords=query, region="wt-wt", safesearch="moderate", max_results=10) # Increased from 3 to 10 for more comprehensive results
        return results
    except Exception as e:
        print("DuckDuckGo text search error:", e)
        return []

def generate_search_context(search_results):
    """
    Generates a combined text string from search results to provide context to Gemini.
    """
    combined_text = ""
    for result in search_results:
        title = result.get("title", "No title")
        snippet = result.get("body", "No snippet available")
        url = result.get("href", "")
        combined_text += f"Title: {title}\nSnippet: {snippet}\nURL: {url}\n\n"
    return combined_text

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    """Extract text from uploaded files (PDF, DOCX, TXT)"""
    filename = secure_filename(file.filename)
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        # Process based on file type
        if file_ext == 'pdf':
            file_stream = io.BytesIO(file.read())
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                text += pdf_reader.pages[page_num].extract_text() + "\n"
            return text
            
        elif file_ext == 'docx':
            file_stream = io.BytesIO(file.read())
            doc = docx.Document(file_stream)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
            
        elif file_ext == 'txt':
            return file.read().decode('utf-8')
            
        return "Unsupported file format"
        
    except Exception as e:
        print(f"Error extracting text from file: {e}")
        return f"Error processing file: {str(e)}"

@app.route("/api/upload", methods=["POST"])
@login_required
def upload_document():
    """Handle document uploads and extract text"""
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400
        
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
        
    if not allowed_file(file.filename):
        return jsonify({"error": "File type not allowed"}), 400
    
    try:
        # Extract text from the file
        extracted_text = extract_text_from_file(file)
        
        if not extracted_text or extracted_text.strip() == "":
            return jsonify({"error": "Could not extract text from file"}), 400
            
        # Get user ID and the upload context
        user_id = current_user.get_id()
        context_id = request.form.get('contextId', secrets.token_hex(8))
        
        # Save text to MongoDB for chat context
        document_data = {
            'user_id': ObjectId(user_id),
            'filename': secure_filename(file.filename),
            'context_id': context_id, 
            'uploaded_at': datetime.utcnow(),
            'text_content': extracted_text,
            'file_type': file.filename.rsplit('.', 1)[1].lower()
        }
        
        mongo.db.document_context.insert_one(document_data)
        
        # Return context ID for the frontend to use
        return jsonify({
            "success": True,
            "context_id": context_id,
            "filename": secure_filename(file.filename),
            "excerpt": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
        })
        
    except Exception as e:
        print(f"Error uploading document: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/documents", methods=["GET"])
@login_required
def get_user_documents():
    """Get list of uploaded documents for the current user"""
    try:
        user_id = current_user.get_id()
        documents = list(mongo.db.document_context.find(
            {'user_id': ObjectId(user_id)},
            {'_id': 1, 'filename': 1, 'context_id': 1, 'uploaded_at': 1, 'file_type': 1}
        ).sort('uploaded_at', -1))
        
        # Convert ObjectId to string for JSON serialization
        for doc in documents:
            doc['_id'] = str(doc['_id'])
            doc['uploaded_at'] = doc['uploaded_at'].isoformat()
            
        return jsonify({"documents": documents})
    except Exception as e:
        print(f"Error retrieving documents: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/documents/<context_id>", methods=["GET"])
@login_required
def get_document_content(context_id):
    """Get content of a specific document"""
    try:
        user_id = current_user.get_id()
        document = mongo.db.document_context.find_one({
            'user_id': ObjectId(user_id),
            'context_id': context_id
        })
        
        if not document:
            return jsonify({"error": "Document not found"}), 404
            
        document['_id'] = str(document['_id'])
        document['uploaded_at'] = document['uploaded_at'].isoformat()
        
        return jsonify({"document": document})
    except Exception as e:
        print(f"Error retrieving document content: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/documents/<context_id>", methods=["DELETE"])
@login_required
def delete_document(context_id):
    """Delete a specific document"""
    try:
        user_id = current_user.get_id()
        result = mongo.db.document_context.delete_one({
            'user_id': ObjectId(user_id),
            'context_id': context_id
        })
        
        if result.deleted_count == 0:
            return jsonify({"error": "Document not found"}), 404
            
        return jsonify({"success": True})
    except Exception as e:
        print(f"Error deleting document: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500

def query_gemini_chat(messages, context=None, deep_analysis=False, use_search=False, document_context_id=None):
    """
    Uses the Gemini API to generate a chat response based on conversation history,
    with optional deep analysis, document context, and DuckDuckGo search integration.
    """
    try:
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        model = genai.GenerativeModel("gemini-2.0-flash") # or "gemini-pro"

        gemini_messages = []
        
        # Add regular context if provided
        if context:
            gemini_messages.append({"role": "user", "parts": [{"text": f"Context for this conversation: {context}"}]})
        
        # Add document context if provided
        if document_context_id:
            document = mongo.db.document_context.find_one({'context_id': document_context_id})
            if document and 'text_content' in document:
                document_text = document['text_content']
                # Truncate if extremely long
                if len(document_text) > 25000:
                    document_text = document_text[:25000] + "... [Document truncated due to length]"
                
                gemini_messages.append({
                    "role": "user", 
                    "parts": [{"text": f"Document context for this conversation (from file: {document['filename']}):\n\n{document_text}"}]
                })

        # Incorporate search results if use_search is True
        search_context_text = ""
        if use_search and messages: # Only search if there are messages (to have a query)
            last_user_message_content = messages[-1]['content']
            search_results = query_duckduckgo_text(last_user_message_content)
            search_context_text = generate_search_context(search_results)
            if search_context_text:
                gemini_messages.append({"role": "user", "parts": [{"text": f"Search Results:\n{search_context_text}"}]}) # Add search results as context

        for msg in messages:
            gemini_messages.append({"role": msg['role'], "parts": [{"text": msg['content']}]})

        prompt_prefix = ""
        if deep_analysis:
            prompt_prefix += (
                "Perform a deep and comprehensive analysis in your response. "
                "Consider multiple perspectives, underlying mechanisms, and potential implications. "
            )

        # Prepend prompt prefix to the last user message (which is the current turn's prompt)
        if gemini_messages and gemini_messages[-1]['role'] == 'user': # Ensure there's a user message
            last_message_content = gemini_messages[-1]['parts'][0]['text']
            gemini_messages[-1]['parts'][0]['text'] = prompt_prefix + last_message_content

        chat = model.start_chat(history=gemini_messages)
        response = chat.send_message(gemini_messages[-1]["parts"][0]["text"]) # Send the (potentially modified) last user message

        return response.text
    except Exception as e:
        print("Gemini API chat error:", e)
        return None

@app.route("/api/chat", methods=["POST"])
def handle_chat():
    data = request.json
    messages = data.get("messages", [])
    context = data.get("context")
    deep_analysis_enabled = data.get("deep_analysis", False)
    use_search_enabled = data.get("use_search", False)
    document_context_id = data.get("document_context_id")

    if not messages:
        return jsonify({"error": "No messages provided"}), 400

    gemini_response_text = query_gemini_chat(
        messages,
        context,
        deep_analysis=deep_analysis_enabled,
        use_search=use_search_enabled,
        document_context_id=document_context_id
    )

    if gemini_response_text is None:
        return jsonify({"error": "Error generating chat response from Gemini"}), 500

    final_response = {
        "response": gemini_response_text,
        "deep_analysis": deep_analysis_enabled,
        "use_search": use_search_enabled,
        "document_processed": bool(document_context_id)
    }
    
    # Log search history for logged-in users
    if current_user.is_authenticated and len(messages) > 0:
        query = messages[-1]['content']
        search_type = "deep" if deep_analysis_enabled else "quick"
        SearchHistory.add_search(mongo.db, current_user.get_id(), query, search_type, document_id=document_context_id)
    
    return jsonify(final_response)

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/api/query", methods=["POST"])
def handle_query():
    data = request.json
    user_query = data.get("query", "")
    deep_analysis = data.get("enable_deep_analysis", False)

    if not user_query:
        return jsonify({"error": "No query provided"}), 400

    search_results = query_duckduckgo_text(user_query)
    generated_answer = generate_answer_from_search(user_query, search_results, deep_analysis)

    if generated_answer is None:
        return jsonify({"error": "Error generating answer from Gemini"}), 500

    final_response = {
        "query": user_query,
        "search_results": search_results,
        "answer": generated_answer,
        "deep_analysis": deep_analysis
    }
    
    # Log search history for logged-in users
    if current_user.is_authenticated:
        search_type = "deep" if deep_analysis else "quick"
        results_count = len(search_results) if search_results else 0
        SearchHistory.add_search(
            mongo.db, 
            current_user.get_id(), 
            user_query, 
            search_type, 
            results_count,
            None,  # No papers for regular searches
            generated_answer,  # Save the answer
            search_results  # Save the search results
        )
    
    return jsonify(final_response)


def generate_answer_from_search(user_query, search_results, deep_analysis=False):
    """
    Generates answer with optional deep analysis mode
    """
    combined_text = ""
    # Remove the [:3] slice to use all search results instead of just the first 3
    for result in search_results:
        title = result.get("title", "No title")
        snippet = result.get("body", "No snippet available")
        url = result.get("href", "")
        combined_text += f"Title: {title}\nSnippet: {snippet}\nURL: {url}\n\n"

    base_prompt = (
        "To generate a response for the following query, follow these steps:\n\n"
        "1. Perform comprehensive analysis of all provided sources\n"
        "2. Identify technical specifications and underlying mechanisms\n"
        "3. Cross-reference with known similar systems\n"
    )

    if deep_analysis:
        base_prompt += (
            "4. Conduct multi-perspective evaluation (technical, social, ethical)\n"
            "5. Generate predictive models for future developments\n"
            "6. Formulate potential failure scenarios and mitigation strategies\n\n"
        )
    else:
        base_prompt += (
            "4. Extract key actionable insights\n"
            "5. Summarize in clear, concise points\n\n"
        )

    prompt = (
        f"{base_prompt}"
        "Now, apply these steps to the following data:\n\n"
        f"{combined_text}\n"
        f"Query: {user_query}\n\n"
        "Provide a comprehensive response:\n"
    )

    return query_gemini(prompt,deep_analysis)


def query_gemini(prompt, deep_analysis=False):  # Add deep_analysis parameter
    try:
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        model = genai.GenerativeModel(
            "gemini-2.0-flash",
            generation_config={
                "max_output_tokens": 2048 if deep_analysis else 1024,
                "temperature": 0.7 if deep_analysis else 0.3
            }
        )
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        print("Gemini API error:", e)
        return None

def query_arxiv(query, max_results=10):  # Increased from 5 to 10
    """
    Query arXiv for academic papers based on the given query.
    Returns a list of paper data.
    """
    try:
        # Construct the API query
        search = arxiv.Search(
            query=query,
            max_results=max_results,
            sort_by=arxiv.SortCriterion.Relevance
        )
        
        # Execute the search and collect results
        print("Executing arXiv search...")
        papers = []
        paper_ids = []  # Keep track of IDs to avoid duplicates
        
        for result in arxiv.Client(page_size=max_results, delay_seconds=1, num_retries=2).results(search):
            try:
                # Skip if we already have this paper
                if result.entry_id in paper_ids:
                    continue
                    
                paper_ids.append(result.entry_id)
                
                # Extract publication date safely
                try:
                    published_date = result.published.strftime("%Y-%m-%d") if result.published else "Unknown"
                    published_year = result.published.year if result.published else None
                except:
                    published_date = "Unknown"
                    published_year = None
                
                # Extract authors safely
                try:
                    authors_list = [{"name": author.name} for author in result.authors] if result.authors else []
                    authors = ", ".join(author.name for author in result.authors) if result.authors else "Unknown"
                except:
                    authors_list = []
                    authors = "Unknown"
                
                # Extract categories
                try:
                    categories = result.categories if hasattr(result, 'categories') else []
                except:
                    categories = []
                
                # Try to extract citation count and references (using pattern matching as a fallback since arXiv doesn't directly provide this)
                citations = []
                references = []
                
                # Get arXiv ID
                arxiv_id = result.entry_id.split("/")[-1] if result.entry_id else "unknown"
                
                # Extract DOI if available
                doi = None
                if hasattr(result, 'doi') and result.doi:
                    doi = result.doi
                elif result.entry_id:
                    # Try to find DOI in the summary
                    doi_match = re.search(r'(doi:|DOI:)\s*([^\s]+)', result.summary or '')
                    if doi_match:
                        doi = doi_match.group(2)
                
                paper = {
                    "title": result.title or "Unknown Title",
                    "authors": authors,
                    "authors_list": authors_list,
                    "summary": result.summary or "No summary available",
                    "pdf_url": result.pdf_url or "#",
                    "published": published_date,
                    "published_year": published_year,
                    "arxiv_id": arxiv_id,
                    "categories": categories,
                    "citations": citations,
                    "references": references,
                    "doi": doi,
                    "url": result.entry_id
                }
                papers.append(paper)
                print(f"Found paper: {paper['title']}")
            except Exception as e:
                print(f"Error processing paper result: {str(e)}")
                continue
        
        print(f"Found {len(papers)} papers")
        # Add simplified citation network data
        try:
            # For each paper, simulate relations to other papers
            # In a real implementation, this would use actual citation data
            for i, paper in enumerate(papers):
                # Randomly select papers that might be related to this one
                for j, other_paper in enumerate(papers):
                    if i != j:
                        # Create artificial citation info for visualization purposes
                        if paper["published_year"] and other_paper["published_year"]:
                            if paper["published_year"] > other_paper["published_year"]:
                                # This paper might cite an older paper
                                paper["references"].append({
                                    "title": other_paper["title"],
                                    "arxiv_id": other_paper["arxiv_id"]
                                })
                                other_paper["citations"].append({
                                    "title": paper["title"],
                                    "arxiv_id": paper["arxiv_id"]
                                })
        except Exception as e:
            print(f"Error building citation relationships: {str(e)}")
        
        # Analyze methodology for each paper
        papers = analyze_papers_methodologies(papers)
            
        return papers
    except Exception as e:
        print(f"arXiv search error: {str(e)}")
        import traceback
        traceback.print_exc()
        return []

def generate_arxiv_context(paper_results):
    """
    Generates a combined text string from arXiv papers to provide context to Gemini.
    """
    combined_text = ""
    for i, paper in enumerate(paper_results, 1):
        combined_text += f"Paper {i}:\n"
        combined_text += f"Title: {paper['title']}\n"
        combined_text += f"Authors: {paper['authors']}\n"
        combined_text += f"Published: {paper['published']}\n"
        combined_text += f"Summary: {paper['summary']}\n"
        combined_text += f"URL: {paper['pdf_url']}\n\n"
    return combined_text

@app.route("/api/deep-research", methods=["POST"])
def handle_deep_research():
    try:
        data = request.json
        query = data.get("query", "")
        max_papers = data.get("max_papers", 3)
        
        # Ensure max_papers is an integer and within reasonable bounds
        try:
            max_papers = int(max_papers)
            max_papers = min(max(max_papers, 1), 10)
        except (ValueError, TypeError):
            max_papers = 3
            
        print(f"Deep research request: query='{query}', max_papers={max_papers}")
        
        if not query:
            return jsonify({"error": "No query provided"}), 400
        
        # Get papers from arXiv
        try:
            paper_results = query_arxiv(query, max_papers)
            print(f"Found {len(paper_results)} papers for query: {query}")
        except Exception as e:
            print(f"Error querying arXiv: {str(e)}")
            return jsonify({"error": f"Error querying arXiv: {str(e)}"}), 500
        
        if not paper_results:
            return jsonify({"error": "No relevant papers found"}), 404
        
        # Generate methodology comparison
        methodology_comparison = compare_methodologies(paper_results)
        
        # Generate context from papers
        arxiv_context = generate_arxiv_context(paper_results)
        
        # Create an enhanced prompt for Gemini
        prompt = (
            f"You are an advanced academic research assistant with expertise in analyzing scientific papers. "
            f"I need you to provide a comprehensive analysis of the following arXiv papers to answer this query: '{query}'\n\n"
            f"Research Question: {query}\n\n"
            f"Here are the relevant papers and their summaries:\n\n{arxiv_context}\n\n"
            f"Please provide a detailed analysis that includes:\n"
            f"1. A comprehensive answer to the research question based on the papers\n"
            f"2. Key findings and methodologies from each relevant paper\n"
            f"3. Any consensus or contradictions among the different papers\n"
            f"4. Limitations of current research on this topic\n"
            f"5. Potential directions for future research\n\n"
            f"When citing papers, use the format 'According to [Authors] in [Title]...'\n"
            f"Your analysis should be well-structured with headings and sections, thorough, and scientifically accurate."
        )
        
        print("Sending enhanced prompt to Gemini for deep analysis")
        
        # Generate answer using Gemini
        try:
            # Use more tokens and higher temperature for academic analysis
            genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
            model = genai.GenerativeModel(
                "gemini-2.0-flash",
                generation_config={
                    "max_output_tokens": 4096,
                    "temperature": 0.7,
                    "top_p": 0.95,
                    "top_k": 40
                }
            )
            
            response = model.generate_content(prompt)
            answer = response.text
            
            if not answer or answer.strip() == "":
                raise Exception("Gemini returned empty response")
                
            print("Successfully received analysis from Gemini")
            
        except Exception as e:
            print(f"Error generating answer from Gemini: {str(e)}")
            return jsonify({"error": f"Error generating answer from Gemini: {str(e)}"}), 500
        
        final_response = {
            "query": query,
            "papers": paper_results,
            "answer": answer,
            "feature": "deep_research",
            "analysis_type": "academic",
            "paper_count": len(paper_results),
            "methodology_comparison": methodology_comparison
        }
        
        # Log search history for logged-in users
        if current_user.is_authenticated:
            SearchHistory.add_search(
                mongo.db, 
                current_user.get_id(), 
                query, 
                "academic", 
                len(paper_results), 
                paper_results,
                answer  # Save the academic analysis/answer
            )
        
        return jsonify(final_response)
    
    except Exception as e:
        print(f"Unexpected error in deep research: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

# Authentication routes
@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('profile'))
        
    error = None
    
    if request.method == 'POST':
        email = request.form.get('email_or_username')
        password = request.form.get('password')
        
        if not email or not password:
            error = 'Email/username and password are required'
            return render_template('login.html', error=error)
            
        try:
            # Find user by email or username
            user_data = mongo.db.users.find_one({
                '$or': [
                    {'email': email},
                    {'username': email}
                ],
                'provider': 'local'
            })
            
            if user_data and check_password_hash(user_data.get('password_hash', ''), password):
                user = User(user_data)
                login_user(user)
                
                # Update last login timestamp
                mongo.db.users.update_one(
                    {'_id': user_data['_id']},
                    {'$set': {'last_login': datetime.utcnow()}}
                )
                
                # Track successful login activity
                track_login_activity(mongo.db, str(user_data['_id']), request, success=True)
                
                # Get next parameter or default to profile
                next_page = request.args.get('next', 'profile')
                return redirect(url_for(next_page))
            else:
                # Track failed login attempt if user exists
                if user_data:
                    track_login_activity(mongo.db, str(user_data['_id']), request, success=False)
                error = 'Invalid email/username or password'
        except Exception as e:
            app.logger.error(f"Error during login: {str(e)}")
            error = 'An error occurred. Please try again later.'
    
    return render_template('login.html', error=error)

@app.route("/api/login-activity")
@login_required
def get_user_login_activity():
    """Get login activity history for the current user"""
    try:
        history = get_login_history(mongo.db, current_user.get_id())
        return jsonify({"success": True, "history": history})
    except Exception as e:
        app.logger.error(f"Error getting login activity: {e}")
        return jsonify({
            "success": False,
            "error": "Failed to retrieve login activity"
        }), 500

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if current_user.is_authenticated:
        return redirect(url_for('profile'))
        
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if not all([username, email, password, confirm_password]):
            flash('All fields are required', 'error')
            return render_template('signup.html', error='All fields are required')
            
        if password != confirm_password:
            flash('Passwords do not match', 'error')
            return render_template('signup.html', error='Passwords do not match')
            
        # Check if username or email already exists
        try:
            existing_user = mongo.db.users.find_one({
                '$or': [
                    {'username': username},
                    {'email': email}
                ]
            })
            
            if existing_user:
                if existing_user.get('username') == username:
                    flash('Username already taken', 'error')
                    return render_template('signup.html', error='Username already taken')
                else:
                    flash('Email already registered', 'error')
                    return render_template('signup.html', error='Email already registered')
                    
            # Hash the password
            password_hash = generate_password_hash(password, method='pbkdf2:sha256')
            
            # Create new user
            new_user = {
                'username': username,
                'email': email,
                'password_hash': password_hash,
                'provider': 'local',
                'provider_id': None,
                'created_at': datetime.now(),
                'picture': None,
                'last_login': datetime.now()
            }
            
            result = mongo.db.users.insert_one(new_user)
            
            if result.inserted_id:
                # Create user object and login
                new_user['_id'] = result.inserted_id
                user = User(new_user)
                login_user(user)
                
                flash('Account created successfully!', 'success')
                return redirect(url_for('profile'))
            else:
                flash('Failed to create account', 'error')
                return render_template('signup.html', error='Failed to create account')
                
        except Exception as e:
            print(f"Error during signup: {str(e)}")
            flash('An error occurred during signup', 'error')
            return render_template('signup.html', error='An error occurred during signup')
    
    return render_template('signup.html')

@app.route("/logout")
@login_required
def logout():
    logout_user()
    session.clear()
    flash("You have been logged out successfully", "success")
    return redirect(url_for("home"))

@app.route("/account/settings")
@login_required
def account_settings():
    return render_template("account_settings.html", user=current_user)

@app.route("/account/update", methods=["POST"])
@login_required
def update_account():
    try:
        data = request.form
        username = data.get("username", "").strip()
        name = data.get("name", "").strip()
        email = data.get("email", "").strip()
        
        # Handle profile picture upload
        picture_url = None
        if 'profile_picture' in request.files:
            file = request.files['profile_picture']
            if file and file.filename:
                try:
                    # Use the avatar system's save_profile_picture function
                    filename = save_profile_picture(file, current_user.username)
                    if filename:
                        picture_url = f"/static/uploads/avatars/{filename}"
                    else:
                        flash("Failed to process the profile picture. Please ensure it's a valid image file (PNG, JPG, or GIF) under 5MB.", "error")
                        return redirect(url_for("account_settings"))
                except Exception as e:
                    app.logger.error(f"Error processing profile picture: {e}")
                    flash("Error processing profile picture. Please try a different image.", "error")
                    return redirect(url_for("account_settings"))
        
        # Update user profile
        success, message = current_user.update_profile(
            mongo.db,
            username=username if username else None,
            name=name if name else None,
            email=email if email else None,
            picture=picture_url
        )
        
        if success:
            flash(message, "success")
        else:
            flash(message, "error")
            
        return redirect(url_for("account_settings"))
        
    except Exception as e:
        app.logger.error(f"Error updating account: {e}")
        flash("An error occurred while updating your account", "error")
        return redirect(url_for("account_settings"))

@app.route("/account/change-password", methods=["POST"])
@login_required
def change_password():
    try:
        data = request.form
        current_password = data.get("current_password", "")
        new_password = data.get("new_password", "")
        confirm_password = data.get("confirm_password", "")
        
        # Basic validation
        if new_password != confirm_password:
            flash("New passwords do not match", "error")
            return redirect(url_for("account_settings"))
            
        if len(new_password) < 8:
            flash("New password must be at least 8 characters long", "error")
            return redirect(url_for("account_settings"))
        
        # Change password
        success, message = current_user.change_password(mongo.db, current_password, new_password)
        
        if success:
            flash(message, "success")
        else:
            flash(message, "error")
            
        return redirect(url_for("account_settings"))
        
    except Exception as e:
        print(f"Error changing password: {e}")
        flash("An error occurred while changing your password", "error")
        return redirect(url_for("account_settings"))

@app.route("/account/delete", methods=["POST"])
@login_required
def delete_account():
    try:
        # Confirm with password for email users
        if current_user.provider == 'email':
            password = request.form.get("password", "")
            if not current_user.check_password(password):
                flash("Incorrect password", "error")
                return redirect(url_for("account_settings"))
        
        # Delete account
        success, message = current_user.delete_account(mongo.db)
        
        if success:
            logout_user()
            session.clear()
            flash(message, "success")
            return redirect(url_for("home"))
        else:
            flash(message, "error")
            return redirect(url_for("account_settings"))
            
    except Exception as e:
        print(f"Error deleting account: {e}")
        flash("An error occurred while deleting your account", "error")
        return redirect(url_for("account_settings"))

# GitHub login route
@app.route("/login/github")
def github_login():
    redirect_uri = url_for("github_callback", _external=True)
    print(f"GitHub callback URL: {redirect_uri}")
    return oauth.github.authorize_redirect(redirect_uri)

@app.route("/login/github/callback")
def github_callback():
    try:
        token = oauth.github.authorize_access_token()
        resp = oauth.github.get("user", token=token)
        user_info = resp.json()
        
        # Get email
        resp = oauth.github.get("user/emails", token=token)
        emails = resp.json()
        email = next((email["email"] for email in emails if email["primary"]), None)
        
        # Check if user exists
        user = User.get_by_provider_id(mongo.db, "github", str(user_info["id"]))
        
        if not user:
            # Create a new user
            user_data = {
                "email": email or "",
                "username": user_info.get("login", ""),
                "name": user_info.get("name", ""),
                "picture": user_info.get("avatar_url", ""),
                "provider": "github",
                "provider_id": str(user_info["id"]),
            }
            user = User(user_data).save(mongo.db)
        else:
            # Update user info
            user.name = user_info.get("name", user.name)
            user.picture = user_info.get("avatar_url", user.picture)
            user.save(mongo.db)
        
        login_user(user)
        
        # Track successful GitHub login
        track_login_activity(mongo.db, user.get_id(), request, success=True)
        
        return redirect(url_for("home"))
    except Exception as e:
        app.logger.error(f"Error in GitHub callback: {e}")
        flash(f"Authentication error: {str(e)}", "error")
        return redirect(url_for("login"))

# User profile and settings
@app.route("/profile")
@login_required
def profile():
    search_history = SearchHistory.get_user_history(mongo.db, current_user.get_id())
    favorites = Favorite.get_favorites(mongo.db, current_user.get_id())
    recommendations = Recommendation.get_recommendations(mongo.db, current_user.get_id())
    
    return render_template(
        "profile.html", 
        user=current_user, 
        search_history=search_history,
        favorites=favorites,
        recommendations=recommendations
    )

# API routes for user data
@app.route("/api/history", methods=["GET"])
@login_required
def get_history():
    history = SearchHistory.get_user_history(mongo.db, current_user.get_id())
    # Convert ObjectId to string for JSON serialization
    for item in history:
        item["_id"] = str(item["_id"])
        item["user_id"] = str(item["user_id"])
        item["timestamp"] = item["timestamp"].isoformat()
    
    return jsonify({"history": history})

@app.route("/api/history/clear", methods=["POST"])
@login_required
def clear_history():
    SearchHistory.clear_history(mongo.db, current_user.get_id())
    return jsonify({"success": True})

@app.route("/api/favorites", methods=["GET"])
@login_required
def get_favorites():
    favorites = Favorite.get_favorites(mongo.db, current_user.get_id())
    # Convert ObjectId to string for JSON serialization
    for item in favorites:
        item["_id"] = str(item["_id"])
        item["user_id"] = str(item["user_id"])
        item["timestamp"] = item["timestamp"].isoformat()
    
    return jsonify({"favorites": favorites})

@app.route("/api/favorites/add", methods=["POST"])
@login_required
def add_favorite():
    data = request.json
    Favorite.add_favorite(
        mongo.db,
        current_user.get_id(),
        data.get("name", "Unnamed"),
        data.get("query", ""),
        data.get("search_type", "quick"),
        data.get("result"),
        data.get("paper")
    )
    return jsonify({"success": True})

@app.route("/api/favorites/remove", methods=["POST"])
@login_required
def remove_favorite():
    data = request.json
    Favorite.remove_favorite(mongo.db, data.get("favorite_id"))
    return jsonify({"success": True})

@app.route("/api/recommendations", methods=["GET"])
@login_required
def get_recommendations():
    recommendations = Recommendation.get_recommendations(mongo.db, current_user.get_id())
    return jsonify({"recommendations": recommendations})

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        
        if not email:
            return render_template('forgot_password.html', error='Email is required')
            
        # Check if user exists
        user = mongo.db.users.find_one({'email': email})
        if not user:
            # Don't reveal that the user doesn't exist for security reasons
            return render_template('forgot_password.html', 
                                  message='If an account with that email exists, a password reset link has been sent.')
        
        # Generate reset token
        reset_token = secrets.token_urlsafe(32)
        expiration = datetime.utcnow() + timedelta(seconds=PASSWORD_RESET_EXPIRATION)
        
        # Store token in database
        mongo.db.password_resets.insert_one({
            'email': email,
            'token': reset_token,
            'expires_at': expiration
        })
        
        # Create reset link
        reset_url = url_for('reset_password', token=reset_token, _external=True)
        
        # Send email
        subject = "Reset Your Sentino Password"
        html_body = f"""
        <html>
            <body>
                <h2>Reset Your Password</h2>
                <p>You've requested to reset your password. Click the link below to set a new password:</p>
                <p><a href="{reset_url}">Reset Password</a></p>
                <p>This link will expire in 1 hour.</p>
                <p>If you didn't request this, please ignore this email.</p>
                <p>Thanks,<br>The Sentino Team</p>
            </body>
        </html>
        """
        
        if send_email(email, subject, html_body):
            return render_template('forgot_password.html', 
                                  message='If an account with that email exists, a password reset link has been sent.')
        else:
            return render_template('forgot_password.html', 
                                  error='Failed to send reset email. Please try again later.')
    
    return render_template('forgot_password.html')

@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    # Check if token exists and is valid
    reset_record = mongo.db.password_resets.find_one({
        'token': token,
        'expires_at': {'$gt': datetime.utcnow()}
    })
    
    if not reset_record:
        flash('Password reset link is invalid or has expired')
        return redirect(url_for('forgot_password'))
    
    if request.method == 'POST':
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if not password or not confirm_password:
            return render_template('reset_password.html', token=token, error='Both fields are required')
            
        if password != confirm_password:
            return render_template('reset_password.html', token=token, error='Passwords do not match')
            
        if len(password) < 8:
            return render_template('reset_password.html', token=token, error='Password must be at least 8 characters')
            
        # Update user's password
        user = mongo.db.users.find_one({'email': reset_record['email']})
        
        if not user:
            return render_template('reset_password.html', token=token, error='User not found')
            
        # Hash new password
        password_hash = generate_password_hash(password, method='pbkdf2:sha256')
        
        # Update user record
        mongo.db.users.update_one(
            {'_id': user['_id']},
            {'$set': {'password_hash': password_hash}}
        )
        
        # Remove used token
        mongo.db.password_resets.delete_one({'token': token})
        
        # Redirect to login page with success message
        flash('Your password has been updated successfully. You can now log in with your new password.')
        return redirect(url_for('login'))
        
    return render_template('reset_password.html', token=token)

@app.route("/login/email", methods=["POST"])
def email_login():
    if current_user.is_authenticated:
        return redirect(url_for('profile'))
    
    email_or_username = request.form.get('email_or_username')
    password = request.form.get('password')
    
    if not email_or_username or not password:
        return redirect(url_for('login', error='Email/username and password are required'))
        
    try:
        # Find user by email or username
        user_data = mongo.db.users.find_one({
            '$or': [
                {'email': email_or_username},
                {'username': email_or_username}
            ],
            'provider': 'local'
        })
        
        if user_data and check_password_hash(user_data.get('password_hash', ''), password):
            user = User(user_data)
            login_user(user)
            
            # Update last login timestamp
            mongo.db.users.update_one(
                {'_id': user_data['_id']},
                {'$set': {'last_login': datetime.utcnow()}}
            )
            
            # Get next parameter or default to profile
            next_page = request.args.get('next', 'profile')
            return redirect(url_for(next_page))
        else:
            return redirect(url_for('login', error='Invalid email/username or password'))
    except Exception as e:
        print(f"Error during email login: {str(e)}")
        return redirect(url_for('login', error='An error occurred. Please try again later.'))

@app.route("/account/update-preferences", methods=["POST"])
@login_required
def update_preferences():
    try:
        data = request.form
        preference_type = data.get("preference_type")
        preferences = {}
        
        # Theme preferences
        if preference_type == "theme":
            preferences = {
                "theme_mode": data.get("theme_mode", "light"),
                "accent_color": data.get("accent_color", "#007bff"),
                "font_size": data.get("font_size", "medium")
            }
        
        # Search preferences
        elif preference_type == "search":
            preferences = {
                "default_search_type": data.get("default_search_type", "quick"),
                "results_per_page": int(data.get("results_per_page", 5)),
                "auto_save_searches": "auto_save_searches" in data
            }
        
        # Privacy preferences
        elif preference_type == "privacy":
            preferences = {
                "history_retention": int(data.get("history_retention", 90)),
                "allow_recommendations": "allow_recommendations" in data,
                "share_usage_data": "share_usage_data" in data
            }
        
        # Notification preferences
        elif preference_type == "notifications":
            preferences = {
                "email_new_recommendations": "email_new_recommendations" in data,
                "email_security_alerts": "email_security_alerts" in data,
                "email_product_updates": "email_product_updates" in data,
                "notification_frequency": data.get("notification_frequency", "daily")
            }
        
        # Update user preferences
        if preferences:
            # Namespace the preferences under their type
            namespaced_prefs = {
                preference_type: preferences
            }
            
            success, message = current_user.update_profile(
                mongo.db,
                preferences=namespaced_prefs
            )
            
            if success:
                flash(f"Your {preference_type} preferences have been updated", "success")
            else:
                flash(message, "error")
        
        return redirect(url_for("account_settings"))
        
    except Exception as e:
        print(f"Error updating preferences: {e}")
        flash("An error occurred while updating your preferences", "error")
        return redirect(url_for("account_settings"))

@app.route("/account/toggle-2fa", methods=["POST"])
@login_required
def toggle_two_factor():
    try:
        data = request.json
        enabled = data.get("enabled", False)
        
        # Update preferences
        preferences = {
            "security": {
                "two_factor_enabled": enabled
            }
        }
        
        success, message = current_user.update_profile(
            mongo.db,
            preferences=preferences
        )
        
        if success:
            return jsonify({
                "success": True,
                "message": "Two-factor authentication " + ("enabled" if enabled else "disabled")
            })
        else:
            return jsonify({
                "success": False,
                "message": message
            }), 400
            
    except Exception as e:
        print(f"Error toggling 2FA: {e}")
        return jsonify({
            "success": False,
            "message": "An error occurred while updating two-factor authentication"
        }), 500

@app.route("/account/setup-2fa", methods=["GET"])
@login_required
def setup_two_factor():
    # In a real implementation, this would generate a QR code for the user's authenticator app
    # For this demo, we'll just show a mock setup page
    return render_template("setup_2fa.html", user=current_user)

@app.route("/account/verify-2fa", methods=["POST"])
@login_required
def verify_two_factor():
    try:
        code = request.form.get("verification_code")
        
        # In a real implementation, this would verify the code against the user's 2FA secret
        # For this demo, we'll just accept any 6-digit code
        if code and len(code) == 6 and code.isdigit():
            # Enable 2FA for the user
            preferences = {
                "security": {
                    "two_factor_enabled": True,
                    "two_factor_verified": True
                }
            }
            
            success, message = current_user.update_profile(
                mongo.db,
                preferences=preferences
            )
            
            if success:
                flash("Two-factor authentication has been enabled for your account", "success")
            else:
                flash(message, "error")
        else:
            flash("Invalid verification code. Please try again.", "error")
            
        return redirect(url_for("account_settings"))
        
    except Exception as e:
        print(f"Error verifying 2FA: {e}")
        flash("An error occurred while verifying two-factor authentication", "error")
        return redirect(url_for("account_settings"))

@app.route("/account/logout-all-devices", methods=["POST"])
@login_required
def logout_all_devices():
    try:
        # Revoke all sessions for the current user
        # In a real implementation, this would revoke all session tokens in the database
        
        # For this implementation, we'll generate a new session token for the current user
        # which will invalidate all other sessions
        current_user.session_token = str(uuid.uuid4())
        success = mongo.db.users.update_one(
            {"_id": current_user._id},
            {"$set": {"session_token": current_user.session_token}}
        ).modified_count > 0
        
        if success:
            # Update the current session with the new token
            session["session_token"] = current_user.session_token
            return jsonify({
                "success": True,
                "message": "You have been logged out from all other devices"
            })
        else:
            return jsonify({
                "success": False,
                "message": "Failed to revoke sessions. Please try again."
            }), 400
            
    except Exception as e:
        print(f"Error logging out from all devices: {e}")
        return jsonify({
            "success": False,
            "message": "An error occurred while logging out from all devices"
        }), 500

@app.route("/account/download-data")
@login_required
def download_user_data():
    try:
        # Get user data
        user_data = {
            "profile": {
                "username": current_user.username,
                "email": current_user.email,
                "name": current_user.name,
                "provider": current_user.provider,
                "created_on": current_user.created_on.isoformat() if hasattr(current_user, 'created_on') and current_user.created_on else None,
                "preferences": current_user.preferences
            },
            "search_history": [],
            "favorites": []
        }
        
        # Get search history
        search_history = mongo.db.search_history.find({"user_id": current_user._id})
        for item in search_history:
            item.pop("_id", None)  # Remove ObjectId which is not JSON serializable
            item.pop("user_id", None)  # Remove user ID reference
            if "timestamp" in item and item["timestamp"]:
                item["timestamp"] = item["timestamp"].isoformat()
            user_data["search_history"].append(item)
            
        # Get favorites
        favorites = mongo.db.favorites.find({"user_id": current_user._id})
        for item in favorites:
            item.pop("_id", None)
            item.pop("user_id", None)
            if "timestamp" in item and item["timestamp"]:
                item["timestamp"] = item["timestamp"].isoformat()
            user_data["favorites"].append(item)
            
        # If user has documents, include their metadata (not content)
        documents = mongo.db.documents.find({"user_id": current_user._id})
        if documents:
            user_data["documents"] = []
            for doc in documents:
                doc.pop("_id", None)
                doc.pop("user_id", None)
                doc.pop("content", None)  # Don't include the actual content
                if "uploaded_on" in doc and doc["uploaded_on"]:
                    doc["uploaded_on"] = doc["uploaded_on"].isoformat()
                user_data["documents"].append(doc)
        
        # Create a JSON file
        json_data = json.dumps(user_data, indent=2)
        
        # Create response with download attachment
        response = make_response(json_data)
        response.headers["Content-Disposition"] = f"attachment; filename=sentino_data_{current_user.username}.json"
        response.headers["Content-Type"] = "application/json"
        
        return response
        
    except Exception as e:
        print(f"Error downloading user data: {e}")
        flash("An error occurred while downloading your data", "error")
        return redirect(url_for("account_settings"))

@app.route("/api/academic/visualizations", methods=["POST"])
def get_academic_visualizations():
    """API endpoint to generate visualizations for academic papers"""
    try:
        data = request.json
        papers = data.get("papers", [])
        viz_type = data.get("type", "network")  # network, timeline, or authors
        
        if not papers:
            return jsonify({"error": "No papers provided"}), 400
            
        # Ensure papers have necessary fields
        for paper in papers:
            if 'title' not in paper:
                return jsonify({"error": "Paper missing title field"}), 400

        if viz_type == "network":
            figure = create_network_visualization(papers)
            return jsonify({"visualization": "network", "data": figure.get('data', []), "layout": figure.get('layout', {})})
        elif viz_type == "timeline":
            figure = create_timeline_visualization(papers)
            return jsonify({"visualization": "timeline", "data": figure.get('data', []), "layout": figure.get('layout', {})})
        elif viz_type == "authors":
            figure = create_author_visualization(papers)
            return jsonify({"visualization": "authors", "data": figure.get('data', []), "layout": figure.get('layout', {})})
        else:
            return jsonify({"error": "Invalid visualization type"}), 400
            
    except Exception as e:
        print(f"Visualization error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

def create_network_visualization(papers):
    """Generate a citation network visualization using networkx and plotly"""
    try:
        import networkx as nx
        import plotly.graph_objects as go
        import random
        
        # Create a graph
        G = nx.DiGraph()
        
        # Generate IDs for papers if not present
        for i, paper in enumerate(papers):
            if 'arxiv_id' not in paper:
                paper['arxiv_id'] = f"paper-{i}"
        
        # Add nodes (papers)
        for paper in papers:
            # Use a shorter version of the title for display
            short_title = paper['title'][:40] + "..." if len(paper['title']) > 40 else paper['title']
            G.add_node(paper['arxiv_id'], title=short_title, year=paper.get('published_year'))
        
        # Simulate citation relationships for visualization
        # In a real system, this would use actual citation data
        paper_ids = [p['arxiv_id'] for p in papers]
        for i, paper in enumerate(papers):
            # Each paper has a chance to cite 1-3 other papers
            num_citations = random.randint(1, min(3, len(papers)-1))
            for _ in range(num_citations):
                cited_id = random.choice(paper_ids)
                if cited_id != paper['arxiv_id']:  # Don't self-cite
                    G.add_edge(paper['arxiv_id'], cited_id)
        
        # Use a spring layout for the graph
        pos = nx.spring_layout(G)
        
        # Create edges trace
        edge_x = []
        edge_y = []
        for edge in G.edges():
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            edge_x.extend([x0, x1, None])
            edge_y.extend([y0, y1, None])
        
        edge_trace = {
            'x': edge_x, 
            'y': edge_y,
            'line': {'width': 1, 'color': '#888'},
            'hoverinfo': 'none',
            'mode': 'lines',
            'type': 'scatter'
        }
        
        # Create nodes trace
        node_x = []
        node_y = []
        for node in G.nodes():
            x, y = pos[node]
            node_x.append(x)
            node_y.append(y)
        
        # Customize node information for hover text
        node_text = []
        for node in G.nodes():
            node_info = G.nodes[node]
            citations = len([edge for edge in G.edges() if edge[1] == node])
            references = len([edge for edge in G.edges() if edge[0] == node])
            node_text.append(f"{node_info['title']}<br>Year: {node_info.get('year', 'Unknown')}<br>Citations: {citations}<br>References: {references}")
        
        # Get degree of each node for coloring
        node_adjacencies = []
        for node, adjacencies in enumerate(G.adjacency()):
            node_adjacencies.append(len(adjacencies[1]))
            
        node_trace = {
            'x': node_x, 
            'y': node_y,
            'mode': 'markers',
            'hoverinfo': 'text',
            'text': node_text,
            'marker': {
                'showscale': True,
                'colorscale': 'YlGnBu',
                'size': 10,
                'color': node_adjacencies,
                'colorbar': {
                    'thickness': 15,
                    'title': 'Node Connections',
                    'xanchor': 'left',
                    'titleside': 'right'
                },
                'line_width': 2
            },
            'type': 'scatter'
        }
        
        # Create the figure
        figure = {
            'data': [edge_trace, node_trace],
            'layout': {
                'title': 'Paper Citation Network',
                'showlegend': False,
                'hovermode': 'closest',
                'margin': {'b': 20, 'l': 5, 'r': 5, 't': 40},
                'xaxis': {'showgrid': False, 'zeroline': False, 'showticklabels': False},
                'yaxis': {'showgrid': False, 'zeroline': False, 'showticklabels': False}
            }
        }
        
        return figure
        
    except Exception as e:
        print(f"Network visualization error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": f"Network visualization error: {str(e)}"}

def create_timeline_visualization(papers):
    """Generate a timeline visualization of papers using plotly"""
    try:
        import plotly.graph_objects as go
        
        # Prepare data for timeline
        timeline_data = []
        for paper in papers:
            year = paper.get('published_year')
            if not year and paper.get('published'):
                # Try to extract year from date string
                match = re.search(r'(\d{4})', paper['published'])
                if match:
                    year = int(match.group(1))
            
            if year:
                timeline_data.append({
                    'title': paper['title'],
                    'year': year,
                    'authors': paper['authors'],
                    'id': paper.get('arxiv_id', 'unknown')
                })
        
        # Sort by year
        timeline_data = sorted(timeline_data, key=lambda x: x['year'])
        
        # Skip if no timeline data
        if not timeline_data:
            return {"error": "No papers with valid publication years"}
        
        # Add papers to timeline
        years = [paper['year'] for paper in timeline_data]
        # Create abbreviated titles for display
        titles = [paper['title'][:20] + '...' if len(paper['title']) > 20 else paper['title'] for paper in timeline_data]
        hover_texts = [f"Title: {paper['title']}<br>Year: {paper['year']}<br>Authors: {paper['authors']}" 
                      for paper in timeline_data]
        
        # Create scatter plot for papers
        scatter = {
            'x': years,
            'y': [1] * len(years),  # All points on same level
            'mode': 'markers+text',
            'marker': {'size': 15, 'color': 'royalblue'},
            'text': titles,
            'textposition': "top center",
            'hoverinfo': 'text',
            'hovertext': hover_texts,
            'type': 'scatter'
        }
        
        # Create the layout
        layout = {
            'title': "Research Timeline",
            'xaxis': {
                'title': "Year",
                'showgrid': True,
                'dtick': 1  # Show each year
            },
            'yaxis': {
                'showticklabels': False,
                'showgrid': False,
                'zeroline': False
            },
            'showlegend': False,
            'hovermode': 'closest'
        }
        
        # Return the figure data
        return {
            'data': [scatter],
            'layout': layout
        }
        
    except Exception as e:
        print(f"Timeline visualization error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": f"Timeline visualization error: {str(e)}"}

def create_author_visualization(papers):
    """Generate a network visualization of author collaborations using networkx and plotly"""
    try:
        import networkx as nx
        import plotly.graph_objects as go
        
        # Create graph
        G = nx.Graph()
        
        # Extract all unique authors
        all_authors = {}
        
        # Add nodes (authors)
        for paper in papers:
            # Process author string or list
            authors = []
            if 'authors_list' in paper and paper['authors_list']:
                authors = paper['authors_list']
            elif 'authors' in paper:
                # Split comma-separated authors
                author_names = paper['authors'].split(',')
                authors = [{'name': name.strip()} for name in author_names]
            
            # Skip if no authors
            if not authors:
                continue
                
            # Add each author as a node
            for author in authors:
                author_name = author.get('name', 'Unknown')
                if author_name not in all_authors:
                    all_authors[author_name] = []
                all_authors[author_name].append(paper['title'])
                G.add_node(author_name, papers=1)
            
            # Add edges between co-authors
            for i, author1 in enumerate(authors):
                author1_name = author1.get('name', 'Unknown')
                for author2 in authors[i+1:]:
                    author2_name = author2.get('name', 'Unknown')
                    
                    # Add edge or increment weight if it exists
                    if G.has_edge(author1_name, author2_name):
                        G[author1_name][author2_name]['weight'] += 1
                    else:
                        G.add_edge(author1_name, author2_name, weight=1)
        
        # Skip if no authors
        if not G.nodes():
            return {"error": "No authors found in papers"}
        
        # Use a spring layout for the graph
        pos = nx.spring_layout(G)
        
        # Create edges trace
        edge_x = []
        edge_y = []
        edge_text = []
        
        for edge in G.edges():
            x0, y0 = pos[edge[0]]
            x1, y1 = pos[edge[1]]
            edge_x.extend([x0, x1, None])
            edge_y.extend([y0, y1, None])
            edge_text.append(f"Co-authored {G[edge[0]][edge[1]]['weight']} papers")
        
        edge_trace = {
            'x': edge_x, 
            'y': edge_y,
            'line': {'width': 1, 'color': '#888'},
            'hoverinfo': 'text',
            'text': edge_text,
            'mode': 'lines',
            'type': 'scatter'
        }
        
        # Create nodes trace
        node_x = []
        node_y = []
        for node in G.nodes():
            x, y = pos[node]
            node_x.append(x)
            node_y.append(y)
        
        # Customize node information for hover text
        node_text = []
        for node in G.nodes():
            papers = all_authors.get(node, [])
            paper_list = "<br>- ".join(papers[:3])
            if len(papers) > 3:
                paper_list += f"<br>...and {len(papers) - 3} more"
            node_text.append(f"Author: {node}<br>Papers: {len(papers)}<br>- {paper_list}")
        
        # Get number of papers for each author for node size/color
        node_papers = []
        for node in G.nodes():
            node_papers.append(len(all_authors.get(node, [])))
            
        node_trace = {
            'x': node_x, 
            'y': node_y,
            'mode': 'markers',
            'hoverinfo': 'text',
            'text': node_text,
            'marker': {
                'showscale': True,
                'colorscale': 'YlGnBu',
                'size': 15,
                'color': node_papers,
                'colorbar': {
                    'thickness': 15,
                    'title': 'Number of Papers',
                    'xanchor': 'left',
                    'titleside': 'right'
                },
                'line_width': 2
            },
            'type': 'scatter'
        }
        
        # Create the figure
        figure = {
            'data': [edge_trace, node_trace],
            'layout': {
                'title': 'Author Collaboration Network',
                'showlegend': False,
                'hovermode': 'closest',
                'margin': {'b': 20, 'l': 5, 'r': 5, 't': 40},
                'xaxis': {'showgrid': False, 'zeroline': False, 'showticklabels': False},
                'yaxis': {'showgrid': False, 'zeroline': False, 'showticklabels': False}
            }
        }
        
        # Return the figure
        return figure
        
    except Exception as e:
        print(f"Author visualization error: {str(e)}")
        import traceback
        traceback.print_exc()
        return {"error": f"Author visualization error: {str(e)}"}

# Add route to view specific search details
@app.route("/api/history/<search_id>", methods=["GET"])
@login_required
def get_search_details(search_id):
    try:
        search = SearchHistory.get_search_by_id(mongo.db, search_id)
        
        if not search:
            return jsonify({"error": "Search not found"}), 404
            
        # Verify the search belongs to the current user
        if str(search.get("user_id")) != current_user.get_id():
            return jsonify({"error": "Unauthorized access"}), 403
            
        # Convert ObjectId to string for JSON serialization
        search["_id"] = str(search["_id"])
        search["user_id"] = str(search["user_id"])
        search["timestamp"] = search["timestamp"].isoformat()
        
        return jsonify({"search": search})
        
    except Exception as e:
        print(f"Error getting search details: {e}")
        return jsonify({"error": f"Error: {str(e)}"}), 500

@app.route('/chat')
@login_required
def chat():
    """Document chat interface where users can upload and chat about documents"""
    return render_template('feature_disabled.html', 
                          feature_name="AI Document Chat", 
                          message="The document chat feature is temporarily disabled for maintenance.", 
                          user_theme=session.get('theme', 'light'))

@app.route('/set_theme', methods=['POST'])
def set_theme():
    """Set the user's theme preference"""
    data = request.json
    if data and 'theme' in data:
        session['theme'] = data['theme']
    return jsonify({"success": True})

# API endpoints for document chat
@app.route('/api/documents')
@login_required
def get_documents():
    """Get all documents for the current user"""
    return jsonify({
        "success": False,
        "message": "Document chat feature is temporarily disabled for maintenance."
    }), 503

@app.route('/api/documents/upload', methods=['POST'])
@login_required
def upload_chat_document():
    """Upload a document for processing"""
    return jsonify({
        "success": False,
        "message": "Document chat feature is temporarily disabled for maintenance."
    }), 503

@app.route('/api/documents/<document_id>', methods=['DELETE'])
@login_required
def delete_chat_document(document_id):
    """Delete a document"""
    return jsonify({
        "success": False,
        "message": "Document chat feature is temporarily disabled for maintenance."
    }), 503

@app.route('/api/chat/<document_id>', methods=['POST'])
@login_required
def chat_with_document(document_id):
    """Chat with a document"""
    return jsonify({
        "success": False,
        "message": "Document chat feature is temporarily disabled for maintenance."
    }), 503

@app.route('/api/chat/<document_id>/history')
@login_required
def get_chat_history(document_id):
    """Get chat history for a document"""
    return jsonify({
        "success": False,
        "message": "Document chat feature is temporarily disabled for maintenance."
    }), 503

def extract_text_from_pdf(filepath):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        
        text = ""
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text
    except Exception as e:
        app.logger.error(f"Error extracting text from PDF: {str(e)}")
        raise e

def extract_text_from_docx(filepath):
    """Extract text from DOCX file"""
    try:
        import docx
        
        doc = docx.Document(filepath)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text
    except Exception as e:
        app.logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise e

def process_document_embeddings(document_id):
    """Process document embeddings for vector search"""
    try:
        document = mongo.db.document_context.find_one({"_id": ObjectId(document_id)})
        if not document:
            app.logger.error(f"Document {document_id} not found for embedding processing")
            return
        
        # In a real implementation, you would:
        # 1. Split the document into chunks
        # 2. Generate embeddings for each chunk
        # 3. Store the embeddings in the database
        
        # For now, we'll just update the document to mark it as processed
        mongo.db.document_context.update_one(
            {"_id": ObjectId(document_id)},
            {"$set": {"embedding_processed": True}}
        )
        
        app.logger.info(f"Document {document_id} embeddings processed successfully")
    
    except Exception as e:
        app.logger.error(f"Error processing document embeddings: {str(e)}")

def generate_document_response(user_query, document_content):
    """Generate AI response based on document content and user query"""
    # In a real implementation, you would:
    # 1. Use embeddings to find relevant chunks of the document
    # 2. Pass the relevant chunks and user query to an LLM API
    # 3. Return the LLM's response
    
    # For demo purposes, we'll just return a simple response
    try:
        # Simple keyword matching for demo
        query_words = set(user_query.lower().split())
        paragraphs = document_content.split('\n\n')
        
        # Find paragraphs that contain query words
        relevant_paragraphs = []
        for paragraph in paragraphs:
            if len(paragraph.strip()) < 10:  # Skip very short paragraphs
                continue
                
            paragraph_words = set(paragraph.lower().split())
            intersection = query_words.intersection(paragraph_words)
            
            if intersection:
                relevant_paragraphs.append({
                    "text": paragraph,
                    "relevance": len(intersection) / len(query_words)
                })
        
        # Sort by relevance
        relevant_paragraphs.sort(key=lambda x: x["relevance"], reverse=True)
        
        # If we found relevant paragraphs
        if relevant_paragraphs:
            top_paragraphs = relevant_paragraphs[:3]
            context = "\n\n".join([p["text"] for p in top_paragraphs])
            
            response = f"Based on the document, I found the following information related to your query:\n\n{context}\n\nIs there anything specific about this you'd like me to explain further?"
        else:
            response = "I couldn't find information directly related to your query in this document. Could you please rephrase your question or ask about a different topic from the document?"
        
        return response
        
    except Exception as e:
        app.logger.error(f"Error generating document response: {str(e)}")
        return "I'm sorry, but I encountered an error while processing your query. Please try again later."

@app.route("/api/methodology-filter", methods=["POST"])
def filter_by_methodology():
    """API endpoint to filter papers by methodology type"""
    try:
        data = request.json
        papers = data.get("papers", [])
        method_type = data.get("methodology_type", "all")
        
        if not papers:
            return jsonify({"error": "No papers provided"}), 400
        
        # If filtering for all, just return all papers
        if method_type == "all":
            return jsonify({"papers": papers})
        
        # Filter papers by methodology type
        filtered_papers = []
        for paper in papers:
            if "methodology" in paper and paper["methodology"]["primary_type"] == method_type:
                filtered_papers.append(paper)
        
        # Generate methodology comparison for the filtered papers
        methodology_comparison = compare_methodologies(filtered_papers)
        
        return jsonify({
            "papers": filtered_papers,
            "methodology_count": len(filtered_papers),
            "methodology_comparison": methodology_comparison
        })
        
    except Exception as e:
        print(f"Error filtering papers by methodology: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/methodology-comparison", methods=["POST"])
def get_methodology_comparison():
    """API endpoint to get methodology comparison for a set of papers"""
    try:
        data = request.json
        papers = data.get("papers", [])
        
        if not papers:
            return jsonify({"error": "No papers provided"}), 400
        
        # Generate methodology comparison
        comparison = compare_methodologies(papers)
        
        return jsonify({
            "comparison": comparison
        })
        
    except Exception as e:
        print(f"Error generating methodology comparison: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

@app.route("/api/analyze-methodology", methods=["POST"])
def handle_methodology_analysis():
    """API endpoint to analyze methodology of a given text"""
    try:
        data = request.json
        text = data.get("text")
        
        if not text:
            return jsonify({"error": "No text provided"}), 400
            
        # Analyze methodology
        methodology_analysis = analyze_methodology(text)
        
        return jsonify({
            "analysis": methodology_analysis,
            "success": True
        })
        
    except Exception as e:
        print(f"Methodology analysis error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Server error: {str(e)}"}), 500

if __name__ == "__main__":
    port = int(os.getenv("PORT", 10000))
    app.run(host='0.0.0.0', debug=True, port=port)